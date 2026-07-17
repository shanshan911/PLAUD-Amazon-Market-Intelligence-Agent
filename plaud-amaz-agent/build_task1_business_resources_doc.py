from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = "Task1_业务方数据资源需求文档.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
HEADER_FILL = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
BORDER = "D9E2EC"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total_twips = sum(int(w * 1440) for w in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def format_table(table, widths, header=True):
    set_table_width(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.05)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor.from_string(INK)
                    set_east_asia_font(run)
            if header and r_idx == 0:
                shade_cell(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        set_paragraph_spacing(paragraph, before=16, after=8, line=1.1)
        size = 16
        color = BLUE
    elif level == 2:
        set_paragraph_spacing(paragraph, before=12, after=6, line=1.1)
        size = 13
        color = BLUE
    else:
        set_paragraph_spacing(paragraph, before=8, after=4, line=1.1)
        size = 12
        color = DARK_BLUE
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(color)
        set_east_asia_font(run)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.1)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(INK)
        set_east_asia_font(run)
        text = text[len(bold_prefix) :]
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, before=0, after=8, line=1.167)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia_font(run)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(paragraph, before=0, after=8, line=1.167)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia_font(run)
    return paragraph


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    format_table(table, [6.5], header=False)
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=2, line=1.1)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia_font(run)
    paragraph = cell.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.1)
    run = paragraph.add_run(body)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    set_east_asia_font(run)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6, line=1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    format_table(table, widths, header=True)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=8, line=1)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Task 1 数据资源需求"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, before=0, after=0, line=1)
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        set_east_asia_font(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = f"生成日期：{date.today().isoformat()}"
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        set_east_asia_font(run)


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=6, line=1.05)
    run = title.add_run("Task 1 业务方数据资源需求文档")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia_font(run)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=12, line=1.1)
    run = subtitle.add_run("适用于欧美日亚马逊站点类目监控、品牌/竞品市占分析与 AI 竞品趋势分析")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_east_asia_font(run)

    add_table(
        doc,
        ["文档项", "说明"],
        [
            ["监控周期", "每周一次"],
            ["监控站点", "US、UK、DE、FR、IT、ES、JP"],
            ["核心任务", "采集卖家精灵类目市场分析报告，计算品牌市占和 AI 竞品市占，并做周环比分析"],
            ["业务用途", "明确业务方需提供的数据、账号、口径和历史资料，保障技术侧稳定实施"],
        ],
        [1.45, 5.05],
    )

    add_callout(
        doc,
        "最低启动条件",
        "业务方至少需提供各站点核心关键词、目标类目路径、PLAUD 品牌别名、竞品品牌清单、AI 筛选规则、卖家精灵账号权限、上一周或历史数据，以及周报输出要求。",
    )

    add_heading(doc, "1. 文档目的", 1)
    add_body(
        doc,
        "本文档用于明确 Task 1 周度监控项目中业务方需提前提供的数据资源、账号环境、统计口径和交付要求。技术侧将基于这些输入完成 Excel 自动解析、指标计算、周环比趋势分析和周报输出。",
    )

    add_heading(doc, "2. 监控范围", 1)
    add_body(doc, "本任务覆盖以下亚马逊站点，业务方需确认各站点均纳入周度监控。")
    add_table(
        doc,
        ["站点", "国家/地区", "是否监控", "备注"],
        [
            ["US", "美国", "是", "需提供关键词与目标类目路径"],
            ["UK", "英国", "是", "需提供关键词与目标类目路径"],
            ["DE", "德国", "是", "需确认德语或英文关键词"],
            ["FR", "法国", "是", "需确认法语或英文关键词"],
            ["IT", "意大利", "是", "需确认意大利语或英文关键词"],
            ["ES", "西班牙", "是", "AI 关键词需同时考虑 AI 与 IA"],
            ["JP", "日本", "是", "需确认日语或英文关键词"],
        ],
        [0.75, 1.25, 1.0, 3.5],
    )

    add_heading(doc, "3. 业务方需提供的数据资源总览", 1)
    add_table(
        doc,
        ["资源类别", "是否必需", "业务方需提供内容", "用途"],
        [
            ["核心搜索关键词", "必需", "7 个站点的搜索关键词及语言口径", "定位亚马逊前台商品与类目入口"],
            ["目标类目路径", "必需", "各站点最终 BSR 类目路径", "校验卖家精灵报告是否来自正确类目"],
            ["PLAUD 品牌别名", "必需", "PLAUD 及其可能出现的所有品牌名写法", "统一识别自营品牌并剔除自营产品"],
            ["竞品品牌清单", "必需", "日常监控竞品品牌及别名", "统计竞品品牌市占与竞品合计"],
            ["AI 筛选规则", "必需", "各站点 AI/IA/本地语言关键词规则", "识别 AI 竞品 ASIN"],
            ["卖家精灵账号与插件", "必需", "账号、插件、导出权限与使用 SOP", "下载类目 Excel 市场分析报告"],
            ["亚马逊访问环境", "必需", "浏览器 Profile、站点语言、邮编和风控处理方式", "保障采集页面一致且可复现"],
            ["历史数据", "必需", "上一周数据，推荐最近 8-12 周", "支持周环比和趋势分析"],
            ["金额与汇率口径", "必需", "币种、汇率来源和换算规则", "统一销售额统计口径"],
            ["输出与分发要求", "必需", "周报格式、接收人、发送时间", "生成符合业务使用场景的交付物"],
        ],
        [1.35, 0.8, 2.1, 2.25],
    )

    add_heading(doc, "4. 详细资源要求", 1)

    add_heading(doc, "4.1 核心搜索关键词", 2)
    add_body(doc, "业务方需提供各站点用于亚马逊前台搜索的核心关键词，并确认是否使用本地语言关键词。")
    add_table(
        doc,
        ["站点", "搜索关键词", "语言要求", "备注"],
        [
            ["US", "Voice Recorders", "英文", "示例，可由业务确认"],
            ["UK", "Voice Recorders", "英文", "示例，可由业务确认"],
            ["DE", "待提供", "德语/英文", "需确认"],
            ["FR", "待提供", "法语/英文", "需确认"],
            ["IT", "待提供", "意大利语/英文", "需确认"],
            ["ES", "待提供", "西班牙语/英文", "需确认"],
            ["JP", "待提供", "日语/英文", "需确认"],
        ],
        [0.75, 2.0, 1.35, 2.4],
    )

    add_heading(doc, "4.2 目标类目路径", 2)
    add_body(doc, "业务方需提供每个站点最终 BSR 类目路径，用于判断采集是否进入正确类目。")
    add_table(
        doc,
        ["站点", "目标类目路径", "是否允许相近类目", "备注"],
        [
            ["US", "Office Products > Electronics > Voice Recorders", "待确认", "示例"],
            ["UK", "待确认", "待确认", ""],
            ["DE", "待确认", "待确认", ""],
            ["FR", "待确认", "待确认", ""],
            ["IT", "待确认", "待确认", ""],
            ["ES", "待确认", "待确认", ""],
            ["JP", "待确认", "待确认", ""],
        ],
        [0.65, 3.15, 1.35, 1.35],
    )

    add_heading(doc, "4.3 PLAUD 品牌标准化信息", 2)
    add_body(doc, "业务方需确认哪些品牌名、店铺名或系列名应统一归为 PLAUD，避免品牌识别漏算。")
    add_table(
        doc,
        ["标准品牌名", "可能出现的别名", "是否纳入 PLAUD", "备注"],
        [
            ["PLAUD", "PLAUD", "是", ""],
            ["PLAUD", "Plaud", "是", ""],
            ["PLAUD", "PLAUD NOTE", "待确认", "需业务确认是否为品牌或系列名"],
            ["PLAUD", "PLAUD.AI", "待确认", "需业务确认"],
        ],
        [1.25, 2.0, 1.35, 1.9],
    )

    add_heading(doc, "4.4 竞品品牌清单", 2)
    add_body(doc, "业务方需提供长期监控的竞品品牌列表，并说明是否纳入竞品合计。")
    add_table(
        doc,
        ["标准竞品品牌名", "品牌别名", "是否核心竞品", "是否纳入竞品合计", "备注"],
        [
            ["待提供", "待提供", "是/否", "是/否", ""],
            ["待提供", "待提供", "是/否", "是/否", ""],
            ["待提供", "待提供", "是/否", "是/否", ""],
        ],
        [1.45, 1.35, 1.1, 1.25, 1.35],
    )

    add_heading(doc, "4.5 AI 竞品关键词规则", 2)
    add_body(doc, "业务方需确认 AI 竞品按标题关键词识别的规则，尤其是 ES 站点是否同时纳入 AI 与 IA。")
    add_table(
        doc,
        ["站点", "建议筛选关键词", "需业务确认事项"],
        [
            ["US", "AI, A.I.", "是否仅按标题识别"],
            ["UK", "AI, A.I.", "是否仅按标题识别"],
            ["DE", "AI, A.I., Künstliche Intelligenz", "是否纳入德语人工智能表达"],
            ["FR", "AI, A.I., Intelligence artificielle", "是否纳入法语人工智能表达"],
            ["IT", "AI, A.I., Intelligenza artificiale", "是否纳入意大利语人工智能表达"],
            ["ES", "AI, IA, A.I., I.A., Inteligencia artificial", "是否排除非人工智能语义下的 IA"],
            ["JP", "AI, A.I., 人工知能", "是否纳入日语人工智能表达"],
        ],
        [0.65, 3.0, 2.85],
    )

    add_heading(doc, "4.6 卖家精灵账号与插件资源", 2)
    add_table(
        doc,
        ["资源项", "是否必需", "业务方需提供内容"],
        [
            ["卖家精灵账号", "必需", "具备市场分析报告导出权限的账号"],
            ["插件安装方式", "必需", "Chrome 插件或指定浏览器插件安装说明"],
            ["账号权限说明", "必需", "是否可查看品牌集中度、商品集中度及导出 Excel"],
            ["登录有效期", "建议提供", "便于安排每周固定采集"],
            ["插件使用 SOP", "建议提供", "包括市场分析、加载全部产品、下载报告等点击路径"],
        ],
        [1.45, 1.0, 4.05],
    )

    add_heading(doc, "4.7 亚马逊访问环境资源", 2)
    add_table(
        doc,
        ["资源项", "是否必需", "说明"],
        [
            ["浏览器环境", "必需", "建议固定 Chrome Profile，保持登录和插件状态稳定"],
            ["亚马逊登录账号", "视情况", "如访问类目、地区设置或插件联动需要，则需提供"],
            ["各站点邮编/地区设置", "建议提供", "避免因地区不同导致商品和类目展示差异"],
            ["语言设置", "建议提供", "确保类目路径、页面字段和导出结果一致"],
            ["验证码/风控处理 SOP", "建议提供", "说明遇到验证码、登录失效或访问限制时如何处理"],
            ["是否允许 RPA 操作", "必需确认", "决定采集流程是否可进入自动化阶段"],
        ],
        [1.55, 1.0, 3.95],
    )

    add_heading(doc, "4.8 历史数据资源", 2)
    add_body(doc, "如无历史数据，第一周只能作为基准周，从第二周开始输出周环比趋势。")
    add_table(
        doc,
        ["数据类型", "最低要求", "推荐要求"],
        [
            ["上周 Excel 原始报告", "必需", "最近 8-12 周"],
            ["上周品牌市占数据", "必需", "最近 8-12 周"],
            ["上周 AI 竞品数据", "必需", "最近 8-12 周"],
            ["历史 ASIN 明细", "建议提供", "最近 8-12 周"],
        ],
        [2.05, 1.6, 2.85],
    )

    add_heading(doc, "4.9 金额与汇率口径", 2)
    add_table(
        doc,
        ["项目", "需确认内容"],
        [
            ["金额币种", "保留各站点本币，还是统一换算 USD 或 RMB"],
            ["汇率来源", "财务汇率、固定汇率、实时汇率或其他指定来源"],
            ["汇率周期", "日汇率、周汇率或月汇率"],
            ["销售额是否含税", "按卖家精灵口径为准，或由业务另行调整"],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "4.10 输出与分发要求", 2)
    add_table(
        doc,
        ["输出项", "需确认内容"],
        [
            ["周报格式", "Excel、Google Sheet、飞书表格、PPT 或 PDF"],
            ["分发方式", "邮件、飞书、钉钉、企业微信或 BI 看板"],
            ["接收人", "业务负责人、运营、管理层或其他固定接收人"],
            ["输出时间", "每周几、几点前完成"],
            ["报告语言", "中文、英文或中英双语"],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "5. 业务方需最终确认的问题", 1)
    questions = [
        "各站点使用哪些搜索关键词？是否使用本地语言关键词？",
        "各站点目标 BSR 类目路径是什么？类目路径变化时是否以最新路径为准？",
        "哪些品牌名、店铺名或系列名应统一归为 PLAUD？",
        "哪些竞品品牌需要长期监控？哪些品牌纳入竞品合计？",
        "AI 竞品是否只按商品标题关键词识别？",
        "ES 站点是否必须同时筛选 AI 和 IA？是否需要排除误匹配？",
        "卖家精灵账号是否具备市场分析报告下载权限？",
        "是否允许使用 RPA 或浏览器自动化辅助采集？",
        "历史数据从哪一周开始提供？是否能提供 8-12 周历史数据？",
        "最终周报以什么格式交付、发送给谁、每周何时发送？",
    ]
    for question in questions:
        add_numbered(doc, question)

    add_heading(doc, "6. 启动前置条件", 1)
    add_body(doc, "技术侧启动稳定实施前，业务方至少需要提供以下资料。")
    prerequisites = [
        "7 个站点的核心搜索关键词。",
        "7 个站点的目标类目路径。",
        "PLAUD 品牌别名表。",
        "竞品品牌清单。",
        "AI 关键词筛选规则。",
        "卖家精灵账号及插件权限。",
        "上一周或历史 Excel 报告。",
        "周报输出格式、输出时间和接收人。",
    ]
    for item in prerequisites:
        add_bullet(doc, item)

    add_heading(doc, "7. 建议交付模板", 1)
    add_body(doc, "业务方可按以下模板一次性补充核心信息。")
    add_table(
        doc,
        ["模块", "业务方填写内容", "负责人", "完成状态"],
        [
            ["站点关键词", "US/UK/DE/FR/IT/ES/JP 对应关键词", "待填写", "待填写"],
            ["目标类目", "各站点 BSR 类目路径", "待填写", "待填写"],
            ["品牌别名", "PLAUD 与竞品品牌别名", "待填写", "待填写"],
            ["AI 规则", "AI/IA/本地语言筛选词", "待填写", "待填写"],
            ["账号环境", "卖家精灵、插件、浏览器、亚马逊环境", "待填写", "待填写"],
            ["历史数据", "上周或 8-12 周原始 Excel 和指标表", "待填写", "待填写"],
            ["周报交付", "格式、接收人、发送时间", "待填写", "待填写"],
        ],
        [1.35, 3.0, 1.05, 1.1],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
